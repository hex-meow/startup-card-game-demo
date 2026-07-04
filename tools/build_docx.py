from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "outputs" / "创业故事卡牌_项目介绍与玩法说明.docx"

FONT = "Microsoft YaHei"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("创业故事卡牌")
    set_run_font(run, size=28, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("项目介绍、流程图与玩法说明")
    set_run_font(run, size=15, color=MUTED)

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(meta)
    widths = [1700, 7660]
    rows = [
        ("版本", "网页端卡牌对战原型"),
        ("核心体验", "从创业路线选择、卡牌构筑、阶段对战到最终发售"),
        ("当前路线", "通用创业赛道 + 初代 iPhone 真实产品故事线"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, LIGHT_BLUE if idx == 0 else WHITE)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(label if idx == 0 else value)
            set_run_font(run, size=10.5, bold=(idx == 0), color=DARK_BLUE if idx == 0 else RGBColor(0, 0, 0))


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.38)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_flow_table(doc, title, steps):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table)
    widths = [1000, 2500, 5860]
    headers = ["顺序", "节点", "说明"]
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(headers[i])
        set_run_font(run, size=10.5, color=DARK_BLUE, bold=True)
    for idx, (node, desc) in enumerate(steps, 1):
        row = table.add_row()
        values = [str(idx), node, desc]
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(values[i])
            set_run_font(run, size=10)


def add_matrix(doc, title, headers, rows, widths=None):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table)
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(headers[i])
        set_run_font(run, size=10.5, bold=True, color=DARK_BLUE)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=9.8)


def build_doc():
    doc = Document()
    style_document(doc)
    add_title_block(doc)

    doc.add_heading("1. 项目概述", level=1)
    add_para(doc, "《创业故事卡牌》是一款网页端卡牌对战原型。玩家通过选择创业路线、构筑牌库、处理阶段关卡和关卡间事件，完成一个项目从立项到发售的全过程。")
    add_bullets(doc, [
        "通用创业路线：银发健康、AI效率工具、绿色消费。玩家通过画像选择决定初始时间、资金和难度。",
        "真实产品故事线：初代 iPhone。作为第四个独立路线选项，直接进入基于 2007 年时间线改编的 Project Purple 到上市流程。",
        "核心表达：用卡牌对战呈现创业中的时间压力、资金约束、团队能力、外部事件和阶段性交付。",
    ])

    add_flow_table(doc, "2. 整体流程图", [
        ("选择路线", "玩家从通用创业赛道或初代 iPhone 真实故事线中选择一条路线。"),
        ("画像或直接开局", "通用路线需要选择核心人群和机会人群；iPhone 线自动带入专属初始牌库。"),
        ("路线地图", "查看 7 个阶段关卡，每个关卡代表一个核心工作目标。"),
        ("阶段对战", "用手牌清空敌人血条，在回合限制内完成核心工作。"),
        ("奖励与事件", "胜利后获得 3 选 1 阶段卡牌，并触发关卡间事件。"),
        ("商店与重组", "用资金购买新卡，重新选择下一关出战牌库。"),
        ("最终发售", "完成最后发售关卡后进入路线收官。"),
    ])

    doc.add_heading("3. 普通创业路线", level=1)
    add_bullets(doc, [
        "银发健康：照护、康复、慢病管理与家庭协同。",
        "AI效率工具：面向企业、创作者与专业岗位的流程提效。",
        "绿色消费：可持续材料、低碳生活方式与新渠道消费。",
    ])
    add_matrix(doc, "画像选择带来的初始变化", ["选择项", "影响内容", "玩法意义"], [
        ("核心人群 3 个", "更强地影响时间、资金、难度", "决定项目最先服务谁，也决定初始风险"),
        ("机会人群 3 个", "以较轻权重影响时间、资金、难度", "代表潜在增长方向，但不会完全主导路线"),
        ("初始牌库 8 张", "从 20 张创业要素牌中选择", "决定第一关战斗风格和资源管理方式"),
    ], widths=[1800, 3300, 4260])

    doc.add_heading("4. 初代 iPhone 真实故事线", level=1)
    add_para(doc, "初代 iPhone 线是一条独立关卡线，不走画像和初始组牌流程。它从 Project Purple 的方向决策开始，经过运营商谈判、OS X 触控移植、Macworld 发布演示、玻璃与续航冲刺、Web App 生态准备，最终到 2007-06-29 发售。")
    add_flow_table(doc, "iPhone 线关卡流程", [
        ("Project Purple立项", "从 iPod、手机和互联网终端的交叉点里，押注全触控设备。"),
        ("运营商谈判", "争取罕见的硬件、软件和营销自主权。"),
        ("OS X触控移植", "把桌面级系统压缩进手机芯片、触控和电池限制里。"),
        ("Macworld发布演示", "2007-01-09，证明它是 iPod、电话和互联网通讯器三合一。"),
        ("玻璃与续航冲刺", "发售前处理抗刮玻璃、续航、射频和量产验证。"),
        ("Web App生态准备", "用 Safari Web 2.0 应用打开第三方生态入口。"),
        ("2007-06-29发售", "完成门店排队、合约激活、媒体评测和首批用户交付。"),
    ])
    add_matrix(doc, "iPhone 线时间线事件", ["事件", "正面影响", "负面压力"], [
        ("2007-01-09：Apple Inc.登场", "品牌叙事升级，阶段奖励资金提高", "公众预期升高，敌人攻击提高"),
        ("2007-01-10：Cisco商标诉讼", "iPhone 名称获得行业关注", "商标纠纷带来法务消耗"),
        ("2007-05-17：FCC批准", "上市路径更明确，难度下降", "规格和节奏被锁定，时间上限下降"),
        ("2007-06-11：WWDC Web 2.0应用", "第三方应用入口打开", "没有原生 SDK 的质疑出现"),
        ("2007-06-18：玻璃与8小时通话", "玻璃与续航消息增强信心", "临近发售的验证压力上升"),
        ("2007-06-28：员工机与首发排队", "首发热度成形", "库存、激活和门店协调压力上升"),
    ], widths=[2600, 3380, 3380])

    doc.add_heading("5. 对战玩法", level=1)
    add_matrix(doc, "核心数值说明", ["数值", "含义", "失败/胜利关联"], [
        ("玩家时间", "玩家血量，敌人攻击会扣减时间", "归零则阶段失败"),
        ("资金", "打出部分卡牌和商店买卡的资源", "资金不足会限制出牌和成长"),
        ("行动力", "每回合打牌资源", "决定单回合处理效率"),
        ("敌人血条", "阶段核心工作剩余量", "清空代表阶段胜利"),
        ("回合限制", "每关必须完成的最大轮数", "超过限制仍未清空血条则失败"),
        ("时间缓冲", "先抵消敌人造成的时间扣减", "提高当前回合生存能力"),
        ("推进加成", "提高后续核心工作推进效果", "增强持续输出"),
        ("阻力/压力", "敌人的防御与攻击成长", "阻力抵消推进，压力提高后续伤害"),
    ], widths=[1700, 4200, 3460])
    add_flow_table(doc, "单局对战流程", [
        ("进入阶段", "洗入出战牌库并抽 5 张手牌。"),
        ("玩家回合", "消耗行动力和资金打出卡牌。"),
        ("处理效果", "推进核心工作，或获得资金、时间、缓冲、抽牌等收益。"),
        ("结束本轮", "弃掉手牌，进入敌人行动。"),
        ("敌人行动", "敌人攻击、增加阻力或增加压力。"),
        ("检查胜负", "敌人血条归零则胜利；玩家时间归零或超回合则失败。"),
    ])

    doc.add_heading("6. 局外成长与商店", level=1)
    add_bullets(doc, [
        "阶段奖励：前 6 个关卡每关都有 3 张独特奖励牌，玩家选择 1 张加入拥有牌库。",
        "关卡间事件：普通路线触发政策、监管、供应、媒体等外部事件；iPhone 线触发 2007 时间线事件。",
        "卡牌商店：关卡之间消耗资金购买未拥有的基础卡。",
        "重组选牌：每完成一个关卡，出战牌库上限 +1；进入关卡至少选择 8 张卡牌。",
        "常驻事件：已触发事件会持续影响后续关卡，并可在路线页和对战页点击查看详情。",
    ])

    doc.add_heading("7. 内容编辑器", level=1)
    add_para(doc, "游戏内提供内容编辑器，可在开局页和路线页打开。它用于快速调参和维护内容。")
    add_matrix(doc, "编辑器范围", ["模块", "可编辑内容", "说明"], [
        ("关卡", "名称、敌人、血量、攻击、回合数、奖励资金、恢复时间、简介", "用于调整阶段难度和叙事"),
        ("事件", "事件名、正面效果、负面效果、常驻数值", "用于调整关卡间环境变化"),
        ("卡牌", "名称、类型、费用、商店价格、描述", "展示信息可改，卡牌逻辑暂时固定"),
    ], widths=[1700, 5000, 2660])

    doc.add_heading("8. 胜利与失败条件", level=1)
    add_bullets(doc, [
        "单关胜利：在回合限制内清空敌人血条。",
        "路线胜利：完成最后发售关卡。",
        "失败条件一：玩家时间归零。",
        "失败条件二：超过关卡回合限制仍未完成核心工作。",
    ])

    doc.add_heading("9. 项目文件与试玩方式", level=1)
    add_matrix(doc, "文件结构", ["文件", "用途"], [
        ("index.html", "页面入口和全局弹窗容器。"),
        ("styles.css", "线稿美术、卡牌、关卡、对战和编辑器样式。"),
        ("app.js", "全部游戏数据、流程状态、卡牌效果、渲染和交互逻辑。"),
        ("项目介绍与玩法说明.md", "Markdown 版说明文档。"),
    ], widths=[2600, 6760])
    add_para(doc, "这是一个静态网页原型。直接打开 index.html 即可运行；如浏览器本地文件策略导致异常，可在 D:\\code 目录下启动本地静态服务后访问。")

    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("创业故事卡牌 · 项目介绍与玩法说明")
        set_run_font(run, size=9, color=MUTED)

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
